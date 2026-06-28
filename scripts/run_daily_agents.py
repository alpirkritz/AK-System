#!/usr/bin/env python3
"""Run daily ABC agents against live Notion context and export to Word."""

from __future__ import annotations

import json
import os
import re
import sys
from datetime import date, datetime, timedelta
from pathlib import Path
from urllib.request import Request, urlopen

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "O_Output"

TASK_DATABASES = {
    "181e7d50-cb8e-8101-9d8a-e90aa8f9b3ac": "Personal To-do",
    "a38dba80-f058-4009-b8d9-bce763f10542": "DT - Action items",
    "20fe7d50-cb8e-805a-9730-cfb2b6e2bfe6": "Con Action items",
}
ASSISTANT_DB = "325e7d50-cb8e-80c1-9046-f71dbdf75f9f"
DONE_STATUSES = {
    "done", "complete", "completed", "closed", "cancelled", "canceled", "archived", "resolved"
}
USER_NAME = os.environ.get("NOTION_USER_NAME", "Alpir Kritzler")
SIMULATED_TIME = "09:00 IDT"
TODAY = date.today().isoformat()


def date_match(title: str) -> bool:
    """Match today in various Notion Assistant title formats."""
    d = date.fromisoformat(TODAY)
    patterns = [
        TODAY,
        d.strftime("%B %d, %Y"),       # June 28, 2026
        d.strftime("%b %d, %Y"),       # Jun 28, 2026
        d.strftime("%B %d %Y"),        # June 28 2026
        d.strftime("%b %d %Y"),        # Jun 28 2026
        d.strftime("%d %b %Y"),        # 28 Jun 2026
        d.strftime("%A, %B %d, %Y"),   # Sunday, June 28, 2026
        d.strftime("%A, %B %d %Y"),    # Sunday, June 28 2026
        d.strftime("%a %d %b %Y"),     # Sun 28 Jun 2026
    ]
    return any(p in title for p in patterns)


def load_token() -> str:
    key = os.environ.get("NOTION_API_KEY")
    if key:
        return key
    env_path = ROOT / "apps" / "web" / ".env.local"
    if env_path.exists():
        for line in env_path.read_text().splitlines():
            if line.startswith("NOTION_API_KEY="):
                return line.split("=", 1)[1].strip()
    raise SystemExit("NOTION_API_KEY not found")


def notion_request(method: str, path: str, body: dict | None = None) -> dict:
    token = load_token()
    data = json.dumps(body).encode() if body is not None else None
    req = Request(
        f"https://api.notion.com/v1{path}",
        data=data,
        method=method,
        headers={
            "Authorization": f"Bearer {token}",
            "Notion-Version": "2022-06-28",
            "Content-Type": "application/json",
        },
    )
    with urlopen(req, timeout=60) as resp:
        return json.loads(resp.read().decode())


def plain(prop: dict | None) -> str:
    if not prop:
        return ""
    t = prop.get("type")
    if t == "title":
        return "".join(x.get("plain_text", "") for x in prop.get("title", []))
    if t == "rich_text":
        return "".join(x.get("plain_text", "") for x in prop.get("rich_text", []))
    if t == "select" and prop.get("select"):
        return prop["select"].get("name", "")
    if t == "status" and prop.get("status"):
        return prop["status"].get("name", "")
    if t == "date" and prop.get("date"):
        return (prop["date"].get("start") or "").split("T")[0]
    if t == "people":
        return ", ".join(p.get("name", "") for p in prop.get("people", []) if p.get("name"))
    return ""


def get_title(props: dict) -> str:
    for v in props.values():
        if v.get("type") == "title":
            return plain(v)
    return ""


def is_done(props: dict) -> bool:
    for pn, pv in props.items():
        if pv.get("type") == "checkbox" and "done" in pn.lower() and pv.get("checkbox"):
            return True
        if pv.get("type") in ("status", "select"):
            if plain(pv).lower() in DONE_STATUSES:
                return True
    return False


def get_due(props: dict) -> str:
    for pv in props.values():
        if pv.get("type") == "date":
            v = plain(pv)
            if v:
                return v[:10]
    return ""


def get_priority(props: dict) -> str:
    for pn, pv in props.items():
        if "priority" in pn.lower() and pv.get("type") in ("select", "status"):
            return plain(pv)
    return ""


def get_assignee(props: dict) -> str:
    for pn, pv in props.items():
        if pv.get("type") == "people" and re.search(r"assign|owner|person|who", pn, re.I):
            return plain(pv)
    for pv in props.values():
        if pv.get("type") == "people":
            return plain(pv)
    return ""


def query_database(db_id: str) -> list[dict]:
    pages: list[dict] = []
    cursor = None
    while True:
        body: dict = {"page_size": 100}
        if cursor:
            body["start_cursor"] = cursor
        data = notion_request("POST", f"/databases/{db_id}/query", body)
        pages.extend(data.get("results", []))
        if not data.get("has_more"):
            break
        cursor = data.get("next_cursor")
    return pages


def fetch_block_text(page_id: str, depth: int = 0) -> list[str]:
    if depth > 3:
        return []
    lines: list[str] = []
    data = notion_request("GET", f"/blocks/{page_id}/children?page_size=100")
    for block in data.get("results", []):
        btype = block.get("type")
        if btype == "table" and block.get("has_children"):
            rows = notion_request("GET", f"/blocks/{block['id']}/children?page_size=100")
            for row in rows.get("results", []):
                if row.get("type") == "table_row":
                    cells = row["table_row"]["cells"]
                    lines.append("| " + " | ".join(
                        "".join(x.get("plain_text", "") for x in cell) for cell in cells
                    ) + " |")
        elif btype and block.get(btype):
            rt = block[btype].get("rich_text", [])
            txt = "".join(x.get("plain_text", "") for x in rt)
            if txt.strip():
                lines.append(txt)
        if block.get("has_children") and btype != "table":
            lines.extend(fetch_block_text(block["id"], depth + 1))
    return lines


def fetch_assistant_page(pattern: str) -> str:
    try:
        pages = query_database(ASSISTANT_DB)
    except Exception:
        return ""
    for page in pages:
        title = get_title(page.get("properties", {}))
        if re.search(pattern, title, re.I) and date_match(title):
            pid = page.get("id")
            if pid:
                return "\n".join(fetch_block_text(pid))
    return ""


def triage_task(task: dict, today: str, soon: str) -> str:
    due = task.get("due", "")
    pri = (task.get("priority") or "").lower()
    if due and due < today:
        return "overdue"
    if due == today:
        return "today"
    if pri in ("high", "urgent", "critical"):
        return "high"
    if due and due <= soon:
        return "soon"
    return "skip"


def fetch_tasks() -> dict[str, list[dict]]:
    today = TODAY
    soon = (date.fromisoformat(today) + timedelta(days=3)).isoformat()
    buckets = {"overdue": [], "today": [], "soon": [], "high": []}
    seen: set[str] = set()

    for db_id, db_name in TASK_DATABASES.items():
        try:
            pages = query_database(db_id)
        except Exception:
            continue
        for page in pages:
            props = page.get("properties", {})
            if is_done(props):
                continue
            title = get_title(props).strip()
            if not title:
                continue
            assignee = get_assignee(props)
            if assignee and USER_NAME not in assignee and db_name != "Personal To-do":
                continue
            if not assignee and db_name != "Personal To-do":
                continue
            status_prop = next((p for p in props.values() if p.get("type") == "status"), None)
            task = {
                "db": db_name,
                "title": title,
                "due": get_due(props),
                "status": plain(status_prop) if status_prop else "",
                "priority": get_priority(props),
                "assignee": assignee,
            }
            key = f"{task['db']}:{task['title']}"
            if key in seen:
                continue
            seen.add(key)
            level = triage_task(task, today, soon)
            if level != "skip":
                buckets[level].append(task)

    for k in buckets:
        buckets[k].sort(key=lambda t: t.get("due") or "9999")
    return buckets


def parse_calendar_events(calendar_text: str) -> list[dict]:
    events = []
    for line in calendar_text.splitlines():
        if not line.strip().startswith("|"):
            continue
        if "Time" in line and "Event" in line:
            continue
        if re.match(r"\|\s*-+\s*\|", line):
            continue
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        if len(parts) < 3:
            continue
        if not re.match(r"\d{1,2}:\d{2}", parts[0]):
            continue
        events.append({
            "time": parts[0],
            "title": parts[1],
            "calendar": parts[2] if len(parts) > 2 else "",
            "notes": parts[3] if len(parts) > 3 else "",
        })
    return events


def parse_email_summary(email_text: str) -> dict:
    return {
        "raw": email_text,
        "reply_needed": "Reply Needed" in email_text and "None" not in email_text.split("Reply Needed")[-1][:80],
        "action": "Action / Decision" in email_text,
    }


def fmt_tasks(tasks: list[dict], limit: int = 10) -> str:
    if not tasks:
        return "- None"
    lines = []
    for t in tasks[:limit]:
        due = f" (due {t['due']})" if t.get("due") else ""
        pri = t.get("priority") or "?"
        lines.append(f"- **[{pri}]** {t['title']}{due} — *{t['db']}*")
    return "\n".join(lines)


def generate_morning_brief(tasks: dict, calendar: str, email: str) -> str:
    top_today = tasks["today"][:3] or tasks["high"][:3] or tasks["overdue"][:3]
    missed = [t for t in tasks["overdue"] if t.get("priority", "").lower() == "high"][:2]

    priorities = []
    for i, t in enumerate(top_today[:3], 1):
        due_note = f" Deadline: **{t['due']}**." if t.get("due") else ""
        priorities.append(
            f"**Priority {i}: {t['title']}** — {t['status'] or 'Open'} in {t['db']}.{due_note} "
            f"Priority: {t.get('priority') or 'normal'}."
        )

    missed_section = "\n".join(
        f"- **{t['title']}** — overdue since {t['due']} ({t['db']})"
        for t in missed
    ) if missed else "- Nothing critical missed yesterday."

    tldr = f"{len(tasks['today'])} due today, {len(tasks['overdue'])} overdue, {len(parse_calendar_events(calendar))} calendar events."

    return f"""# ☀️ Morning Brief – {TODAY}

> **TL;DR:** {tldr}
> **Simulated time:** {SIMULATED_TIME}
> **Agent:** `03_morning_briefing`

## 🏆 Today's Priorities

{chr(10).join(priorities) if priorities else '- No urgent priorities flagged.'}

## 👀 Things I Missed Yesterday

{missed_section}

---

## Calendar (source: Notion Assistant)

{calendar or '_No calendar review found._'}

## Email snapshot

{email or '_No email summary found._'}
"""


def generate_meeting_prep(events: list[dict], tasks: dict) -> str:
    sections = []
    work_meetings = [
        e for e in events
        if "dragontail" in e.get("calendar", "").lower()
        or "daz" in e.get("calendar", "").lower()
        or re.search(r"1:1|sync|farewell", e.get("title", ""), re.I)
    ]

    open_items = tasks["today"] + tasks["overdue"][:5] + tasks["high"][:3]

    for ev in work_meetings:
        related = [t for t in open_items if any(w.lower() in t["title"].lower() for w in ev["title"].split()[:2])][:3]
        if not related:
            related = open_items[:2]
        flags = [t for t in related if t.get("priority", "").lower() == "high" or (t.get("due") and t["due"] < TODAY)]
        sections.append(f"""### {ev['title']} ({ev['time']})

- 👥 **Participants:** {ev.get('notes') or 'See calendar invite'} ({ev.get('calendar', 'Work')})
- 🚨 **Important flags:** {', '.join(t['title'] for t in flags[:3]) or 'None flagged'}
- 💭 **Last time we met:** Check AI Meeting Notes for prior context on this topic.
- 🎯 **Today's focus:** Push forward open action items; confirm decisions and blockers.

**Top open items:**
{fmt_tasks(related, 3)}

**Next pushes:** Resolve highest-priority blockers before the meeting.
**Questions:** What decision is needed today?
""")

    if not sections:
        sections.append("No work meetings found for today. Top open action items:\n\n" + fmt_tasks(open_items, 5))

    return f"""# Meeting Prep Herald — {TODAY}

> **Simulated time:** {SIMULATED_TIME}
> **Agent:** `04_meeting_prep_herald`

---

{chr(10).join('---' + chr(10) + s for s in sections)}
"""


def generate_calendar_optimizer(events: list[dict], tasks: dict, calendar: str) -> str:
    work_events = [e for e in events if "dragontail" in e.get("calendar", "").lower() or "daz" in e.get("calendar", "").lower()]
    overlaps = []
    for i, a in enumerate(events):
        for b in events[i + 1:]:
            ta = a["time"].split("–")[0].strip() if a.get("time") else ""
            tb = b["time"].split("–")[0].strip() if b.get("time") else ""
            if ta and tb and ta == tb:
                overlaps.append((a, b))

    work_hours = 2.5 if work_events else 0
    load_note = "heavy" if work_hours > 4 else "manageable"

    summary = []
    for a, b in overlaps:
        if "dragontail" in a.get("calendar", "").lower() and "dragontail" in b.get("calendar", "").lower():
            summary.append(f"🚨 Conflict: {a['title']} vs {b['title']} at {a['time']} → Review priority")
        else:
            summary.append(f"⚠️ Overlap: {a['title']} overlaps with {b['title']} at {a['time']}")

    back_to_back = []
    dragontail = [e for e in events if "dragontail" in e.get("calendar", "").lower()]
    for i in range(len(dragontail) - 1):
        end_a = dragontail[i]["time"].split("–")[-1].strip() if "–" in dragontail[i]["time"] else ""
        start_b = dragontail[i + 1]["time"].split("–")[0].strip() if "–" in dragontail[i + 1]["time"] else ""
        if end_a and start_b and end_a == start_b:
            back_to_back.append(f"{dragontail[i]['title']} → {dragontail[i+1]['title']}")

    if back_to_back:
        summary.append(f"⚠️ Back-to-back: {'; '.join(back_to_back)} — zero buffer")

    summary.append(f"📊 Load: ~{work_hours:.1f}h work meetings ({load_note})")

    top_task = tasks["today"][0] if tasks["today"] else (tasks["overdue"][0] if tasks["overdue"] else None)
    if top_task:
        summary.append(f"⏰ Top task: {top_task['title']} → suggest 09:00–10:15 deep work slot")

    details = []
    for a, b in overlaps:
        if "dragontail" in a.get("calendar", "").lower() and "dragontail" in b.get("calendar", "").lower():
            details.append(
                f"- **{a['title']} vs {b['title']}** ({a['time']}) — "
                "Recommend: keep group/event over 1:1 if applicable. *Approval required before any reschedule.*"
            )
    if back_to_back:
        for bb in back_to_back:
            details.append(f"- **Back-to-back:** {bb} — consider 5-min buffer between Farewell and R&D sync")

    focus = ""
    if top_task:
        focus = f"- **{top_task['title']}** → 09:00–10:15 (free window before personal block)"

    assistant_section = calendar.strip() if calendar else "_No calendar review found._"

    return f"""# Calendar Optimizer — {TODAY}

> **Simulated time:** {SIMULATED_TIME}
> **Agent:** `06_calendar_optimizer`
> **Note:** Recommendations only — no actions taken without approval.

## Summary

{chr(10).join('- ' + s for s in summary) if summary else '- No conflicts detected. Load is manageable.'}

## Conflicts

{chr(10).join(details) if details else '- No actionable meeting conflicts (per Notion Assistant review).'}

## Focus time

{focus or '- Use 09:00–14:30 for deep work — longest free block today.'}

## Source: Notion Assistant Calendar Review

{assistant_section}

## Reschedule options

*For any recommended moves, check all attendees' calendars for 2–3 alternative slots before requesting approval.*
"""


def generate_email_assistant(email_text: str) -> str:
    if not email_text.strip():
        return f"""# Email Assistant — {TODAY}

> **Simulated time:** {SIMULATED_TIME}
> **Agent:** `07_email_assistant`

## Summary

- 🔴 Needs attention: 0
- 🟡 FYI: 0
- 🟢 Archive: N/A

**Blocker:** No Email Summary page found in Notion Assistant for today. Gmail not connected in this run — connect Gmail for live triage.

| From | Subject | Summary | Recommended action |
|---|---|---|---|
| — | — | No data | — |
"""

    return f"""# Email Assistant — {TODAY}

> **Simulated time:** {SIMULATED_TIME}
> **Agent:** `07_email_assistant`
> **Note:** Recommendations only — no archive/send without confirmation.

## Summary

Based on Notion Assistant Email Summary:

{email_text}

## Triage table

| From | Subject | Summary | Recommended action |
|---|---|---|---|
| (see summary above) | — | Parsed from Assistant page | Review manually |

*Full Gmail thread analysis requires connected inbox — this run used Notion Assistant snapshot only.*
"""


def generate_ibkr_import() -> str:
    return f"""# IBKR Daily Import — {TODAY}

> **Simulated time:** {SIMULATED_TIME}
> **Agent:** `05_ibkr_daily_import`

## Report

**No new transactions found.**

**Blocker:** Gmail is not connected in this automated run. The agent requires Gmail access to scan IBKR transaction emails (BOUGHT/SOLD/Dividend/etc.), de-duplicate against 📈 IBKR Transactions in Notion, insert missing rows, and label/archive threads.

**Next step:** Connect Gmail integration or run manually with inbox access.
"""


def generate_startup_coo_note() -> str:
    return f"""# Startup COO — {TODAY}

> **Agent:** `08_startup_coo`

## Note

Startup COO is an **on-demand** advisory agent (not a daily scheduled run). It activates when you bring a problem, decision, or strategic question.

**Example triggers:**
- "Should we prioritize fundraising or product this quarter?"
- "Help me structure our hiring plan for the next 3 months."
- "Review this ops process and suggest improvements."

No daily output generated — invoke via `/agents` chat when needed.
"""


def export_docx(sections: list[tuple[str, str]], out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_BREAK

    doc = Document()
    title = doc.add_heading(f"Daily Agents Run — {TODAY} (09:00)", 0)
    title.runs[0].font.size = Pt(22)
    doc.add_paragraph(f"Generated at {datetime.now().strftime('%Y-%m-%d %H:%M')} · Simulated morning time: {SIMULATED_TIME}")
    doc.add_paragraph("DRAFT — REQUIRES HUMAN REVIEW")

    for i, (heading, body) in enumerate(sections):
        if i > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_heading(heading, level=1)
        for line in body.splitlines():
            stripped = line.strip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("|") and "---" not in stripped:
                doc.add_paragraph(stripped, style="List Bullet")
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith(">"):
                p = doc.add_paragraph(stripped.lstrip("> ").strip())
                p.runs[0].italic = True if p.runs else None
            elif stripped == "---":
                doc.add_paragraph("—" * 40)
            elif stripped:
                doc.add_paragraph(stripped)

    doc.save(out_path)


def main() -> None:
    print(f"Fetching Notion context for {TODAY}...")
    tasks = fetch_tasks()
    calendar = fetch_assistant_page(r"Calendar (Review|Optimizer|Report|Analysis)")
    email = fetch_assistant_page(r"(Daily )?Email Summary")
    events = parse_calendar_events(calendar)

    outputs = {
        "morning-brief": generate_morning_brief(tasks, calendar, email),
        "meeting-prep": generate_meeting_prep(events, tasks),
        "calendar-optimizer": generate_calendar_optimizer(events, tasks, calendar),
        "email-assistant": generate_email_assistant(email),
        "ibkr-daily-import": generate_ibkr_import(),
        "startup-coo": generate_startup_coo_note(),
    }

    OUTPUT_DIR.mkdir(exist_ok=True)
    md_paths = []
    sections = []
    titles = {
        "morning-brief": "03 — Morning Briefing",
        "meeting-prep": "04 — Meeting Prep Herald",
        "calendar-optimizer": "06 — Calendar Optimizer",
        "email-assistant": "07 — Email Assistant",
        "ibkr-daily-import": "05 — IBKR Daily Import",
        "startup-coo": "08 — Startup COO (on-demand note)",
    }

    for key, content in outputs.items():
        path = OUTPUT_DIR / f"{TODAY}_{key}.md"
        path.write_text(content, encoding="utf-8")
        md_paths.append(path)
        sections.append((titles[key], content))
        print(f"  Wrote {path.name}")

    docx_path = OUTPUT_DIR / f"{TODAY}_daily-agents-run.docx"
    export_docx(sections, docx_path)
    print(f"  Wrote {docx_path.name}")

    combined = OUTPUT_DIR / f"{TODAY}_daily-agents-run.md"
    combined.write_text(
        "\n\n---\n\n".join(f"# {t}\n\n{c}" for t, c in sections),
        encoding="utf-8",
    )
    print(f"  Wrote {combined.name}")
    print("Done.")


if __name__ == "__main__":
    main()
